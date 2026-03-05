import streamlit as st
from groq import Groq
from docx import Document
import json

st.set_page_config(page_title="Consulting AI Intelligence Engine")

st.title("Consulting AI Intelligence Engine")

query = st.text_area("Paste Client Query")

client = Groq(api_key=st.secrets["GROQ_API_KEY"])

CURRENT_YEAR = 2025
FORECAST_YEAR = 2030

# ------------------------------------------------
# CONSULTING FRAMEWORK LIBRARY
# ------------------------------------------------

FRAMEWORK_LIBRARY = {

"market_opportunity":[
"Industry Overview",
"Industry Value Chain",
"Product Definition",
"Technology Architecture",
"Application Landscape",
"Material Substitution Analysis",
"Technology Roadmap",
"Regulatory Landscape",
"Competitive Landscape",
"Regional Market Analysis",
"Market Size Analysis",
"Revenue Pools by Application",
"High Growth Segments",
"Strategic Opportunity Matrix"
],

"technology_landscape":[
"Industry Overview",
"Technology Architecture",
"Material Chemistry",
"Key Applications",
"Emerging Use Cases",
"Innovation Trends",
"Patent Landscape",
"Key Technology Players"
],

"market_entry":[
"Industry Overview",
"Market Structure",
"Customer Segments",
"Competitive Landscape",
"Regulatory Barriers",
"Entry Strategy Options",
"Partnership Opportunities",
"Implementation Roadmap"
]

}

# ------------------------------------------------
# QUERY PARSER
# ------------------------------------------------

def parse_query(query):

    prompt=f"""
Extract structured information from the client query.

Return JSON with fields:

Industry
Product
Applications
Geography
Project_Type
Keywords

If geography not mentioned assume Global.

Query:
{query}
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        temperature=0.2
    )

    return response.choices[0].message.content


# ------------------------------------------------
# INDUSTRY DETECTION
# ------------------------------------------------

def detect_industry(query):

    prompt=f"""
Identify the industry context of this query.

Query:
{query}

Return a short explanation.
"""

    response=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return response.choices[0].message.content


# ------------------------------------------------
# PROJECT TYPE DETECTION
# ------------------------------------------------

def detect_project_type(query):

    prompt=f"""
Classify the consulting project type.

Options:

market_opportunity
technology_landscape
market_entry

Query:
{query}

Return only the classification.
"""

    response=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return response.choices[0].message.content.strip()


# ------------------------------------------------
# BATTERY / SYSTEM ARCHITECTURE
# ------------------------------------------------

def architecture_mapping(query):

    prompt=f"""
Map the system architecture relevant to this query.

For example for EV batteries include:

Cell
Module
Pack
Thermal Management
Battery Electronics

Explain where the product fits in the architecture.

Query:
{query}
"""

    response=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return response.choices[0].message.content


# ------------------------------------------------
# MATERIAL SUBSTITUTION ANALYSIS
# ------------------------------------------------

def material_comparison(query):

    prompt=f"""
Compare the product in this query with alternative materials.

Include:

Performance
Thermal properties
Cost
Advantages
Limitations

Query:
{query}
"""

    response=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return response.choices[0].message.content


# ------------------------------------------------
# APPLICATION REVENUE POOLS
# ------------------------------------------------

def revenue_pools(query):

    prompt=f"""
Identify revenue pools for the product in this market.

Break down by application.

Include estimated opportunity and growth drivers.

Query:
{query}
"""

    response=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return response.choices[0].message.content


# ------------------------------------------------
# KEY PLAYERS
# ------------------------------------------------

def detect_players(query):

    prompt=f"""
Identify leading companies in this market.

Return 8-10 companies with short description.

Query:
{query}
"""

    response=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return response.choices[0].message.content


# ------------------------------------------------
# MODULE GENERATOR
# ------------------------------------------------

def generate_modules(query,framework):

    prompt=f"""
Create consulting research modules.

Sections:
{framework}

For each section include:

Objective
Key Questions
Analysis Required
Deliverables

Include market year {CURRENT_YEAR} and forecast {FORECAST_YEAR} where relevant.

Query:
{query}
"""

    response=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        temperature=0.3
    )

    return response.choices[0].message.content


# ------------------------------------------------
# WORD REPORT
# ------------------------------------------------

def create_word(query,industry,architecture,materials,revenue,players,modules):

    doc=Document()

    doc.add_heading("Consulting Research Framework",level=1)

    doc.add_heading("Client Query",level=2)
    doc.add_paragraph(query)

    doc.add_heading("Industry Context",level=2)
    doc.add_paragraph(industry)

    doc.add_heading("System Architecture",level=2)
    doc.add_paragraph(architecture)

    doc.add_heading("Material Substitution Analysis",level=2)
    doc.add_paragraph(materials)

    doc.add_heading("Revenue Pools",level=2)
    doc.add_paragraph(revenue)

    doc.add_heading("Key Companies",level=2)
    doc.add_paragraph(players)

    doc.add_heading("Consulting Modules",level=2)
    doc.add_paragraph(modules)

    doc.save("consulting_framework.docx")


# ------------------------------------------------
# MAIN ENGINE
# ------------------------------------------------

if st.button("Generate Consulting Framework"):

    st.subheader("Query Analysis")

    parsed=parse_query(query)
    st.write(parsed)

    st.subheader("Industry Detection")

    industry=detect_industry(query)
    st.write(industry)

    project_type=detect_project_type(query)

    if project_type not in FRAMEWORK_LIBRARY:
        project_type="market_opportunity"

    framework=FRAMEWORK_LIBRARY[project_type]

    st.subheader("Architecture Mapping")

    architecture=architecture_mapping(query)
    st.write(architecture)

    st.subheader("Material Substitution")

    materials=material_comparison(query)
    st.write(materials)

    st.subheader("Revenue Pools")

    revenue=revenue_pools(query)
    st.write(revenue)

    st.subheader("Key Companies")

    players=detect_players(query)
    st.write(players)

    st.subheader("Consulting Modules")

    modules=generate_modules(query,framework)
    st.write(modules)

    create_word(query,industry,architecture,materials,revenue,players,modules)

    with open("consulting_framework.docx","rb") as f:

        st.download_button(
            label="Download Word Report",
            data=f,
            file_name="consulting_framework.docx"
        )
